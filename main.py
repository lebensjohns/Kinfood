'''
Sources:
https://pythonprogramming.net/flask-send-file-tutorial/
https://python-docx.readthedocs.io/en/latest/
'''
import datetime

from flask import Flask, render_template
from google.cloud import datastore
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

datastore_client = datastore.Client()

document = Document()

def create_packing_slips():
    doc = docx.
    document.add_heading('Document Title', 0)
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    document.add_page_break()
    return

@app.route('/')
def root():

    # Run api call

    return render_template(
        'index.html')

@app.route('/return-files')
def return_files():

    # document.save('demo.docx')

    return document


if __name__ == '__main__':
    # This is used when running locally only. When deploying to Google App
    # Engine, a webserver process such as Gunicorn will serve the app. This
    # can be configured by adding an `entrypoint` to app.yaml.
    # Flask's development server will automatically serve static files in
    # the "static" directory. See:
    # http://flask.pocoo.org/docs/1.0/quickstart/#static-files. Once deployed,
    # App Engine itself will serve those files as configured in app.yaml.
    app.run(host='127.0.0.1', port=8080, debug=True)